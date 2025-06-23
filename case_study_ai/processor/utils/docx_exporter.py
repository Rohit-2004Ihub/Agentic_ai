from docx import Document

def export_case_to_docx(data, filename="output_case_study.docx"):
    doc = Document()
    doc.add_heading('Project Case Study Report', 0)

    doc.add_heading('Structured Summary', level=1)
    for key, val in data["structured"].items():
        doc.add_paragraph(f"{key.title()}: {val}")

    doc.add_heading('Narrative Case Study', level=1)
    doc.add_paragraph(data["case_study"])

    doc.add_heading('Visual Suggestions', level=1)
    doc.add_paragraph(data["visuals"])

    doc.add_heading('Pitch Feedback', level=1)
    doc.add_paragraph(f"Score: {data['pitch_feedback']['Score']}")
    doc.add_paragraph(f"Strengths: {data['pitch_feedback']['Strengths']}")
    doc.add_paragraph(f"Suggestions: {data['pitch_feedback']['Improvement Suggestions']}")

    doc.save(filename)
